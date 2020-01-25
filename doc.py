from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
#from docx.enum.style import WD_STYLE_TYPE

doc=Document()
doc.add_heading(text='poiuy',level=0)

p=doc.add_paragraph('A plain paragraph haaving some')
p.add_run(' bold').bold=True
p.add_run(' and some')
p.add_run(' italic').italic=True

doc.add_heading('Heading level 1',level=1)
doc.add_paragraph('Intense quote',style='Intense Quote')

doc.add_paragraph('first item in unordered list',style='List Bullet')
doc.add_paragraph('first item in ordered list',style='List Number')

doc.add_picture('2020.png',width=Inches(1.25))

records=((3,'101','Spam'),(7,'422','Eggs'),(4,'631','Spam,spam,eggs,and spam'))

table=doc.add_table(rows=1,cols=3)
hdr_cells=table.rows[0].cells
hdr_cells[0].text='Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

doc.add_page_break()

doc.save('new.docx')
